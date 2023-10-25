from docx import Document

doc = Document('/Users/aleksejdegtarev/Desktop/ecp/1.docx')

text = []
for i in doc.paragraphs:
    text.append(i.text)

date_start = text[4].split('\n')[1].split(': ')[1].strip()
time_start = text[4].split('\n')[2].split(': ')[1].strip()

date_end = text[5].split('\n')[1].split(': ')[1].strip()
time_end = text[5].split('\n')[2].split(': ')[1].strip()
all_days = text[5].split('\n')[3].split(': ')[1].strip()

diagnosis = text[7].split('\n\t')[1].split(': ')[1]
diagnosis_mkb = text[8].split(' \n')[0].split(': ')[1].split()[0]

all_tables = doc.tables

data_tables = {i: None for i in range(len(all_tables))}

for i, table in enumerate(all_tables):
    data_tables[i] = [[] for _ in range(len(table.rows))]
    for j, row in enumerate(table.rows):
        for cell in row.cells:
            data_tables[i][j].append(cell.text)

surname, name, patronymic = data_tables[0][1][1].split(' ')

birthday = data_tables[0][3][1].split('(')[1].strip(')')
